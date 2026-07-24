"""마크다운 -> docx 변환기.

pandoc 이 없는 환경에서 문서를 배포용 Word 파일로 만들기 위한 최소 변환기다.
이 저장소의 문서가 실제로 쓰는 문법만 다룬다:

  # ~ #### 제목 / 문단 / - 목록 / 1. 번호목록 / > 인용
  | 표 |  (구분행 |---| 포함)
  ``` 코드블록 ```  -> 고정폭. mermaid 는 렌더할 수 없으므로 대체 도해를 쓴다.
  **굵게** `코드` [링크](url)  -> 인라인 서식

사용:
  python md2docx.py <입력.md> <출력.docx>
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

# 한글 본문용. Word 는 한글 글자에 별도 폰트 설정(eastAsia)이 필요하다.
BODY_FONT = "맑은 고딕"
MONO_FONT = "D2Coding"      # 없으면 Word 가 대체 고정폭 폰트를 쓴다
MONO_FALLBACK = "Consolas"


def _set_font(run, name: str, size_pt: float = None, bold: bool = None) -> None:
    run.font.name = name
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    # python-docx 는 동아시아 폰트를 별도로 지정해야 한글에 적용된다
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts"
    )
    if rfonts is None:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    from docx.oxml.ns import qn

    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


_INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")


def _add_inline(par, text: str, size_pt: float = 10.0) -> None:
    """**굵게** `코드` [링크](url) 를 서식 있는 run 으로 분해해 넣는다."""
    for part in _INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            _set_font(par.add_run(part[2:-2]), BODY_FONT, size_pt, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = par.add_run(part[1:-1])
            _set_font(run, MONO_FALLBACK, size_pt - 0.5)
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        elif part.startswith("["):
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", part)
            label = m.group(1) if m else part
            _set_font(par.add_run(label), BODY_FONT, size_pt)
        else:
            _set_font(par.add_run(part), BODY_FONT, size_pt)


def _add_code_block(doc, lines: list[str]) -> None:
    """고정폭 블록. 표·도해가 어긋나지 않도록 한 문단에 줄바꿈으로 담는다."""
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Pt(12)
    par.paragraph_format.space_before = Pt(4)
    par.paragraph_format.space_after = Pt(8)
    # 줄 간격을 좁혀 ASCII 도해의 세로선이 이어져 보이게 한다
    par.paragraph_format.line_spacing = 1.0
    for i, line in enumerate(lines):
        if i:
            par.add_run().add_break()
        run = par.add_run(line)
        _set_font(run, MONO_FALLBACK, 8.5)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def _add_table(doc, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            cell = table.cell(ri, ci)
            cell.text = ""
            par = cell.paragraphs[0]
            par.paragraph_format.space_before = Pt(1)
            par.paragraph_format.space_after = Pt(1)
            if ri == 0:
                _set_font(par.add_run(re.sub(r"[*`]", "", cell_text)), BODY_FONT, 9, bold=True)
            else:
                _add_inline(par, cell_text, 9)
    doc.add_paragraph()


def _split_row(line: str) -> list[str]:
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def convert(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    # 기본 스타일
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # --- 코드블록 ---
        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            block: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1
            if lang == "mermaid":
                # Word 는 mermaid 를 렌더하지 못한다. 코드를 그대로 남기면
                # 읽을 수 없으므로 대체 안내를 넣는다(도해는 별도 치환됨).
                par = doc.add_paragraph()
                run = par.add_run("[다이어그램 — 원본 마크다운 참조]")
                _set_font(run, BODY_FONT, 9)
                run.italic = True
            else:
                _add_code_block(doc, block)
            continue

        # --- 표 ---
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(
            r"^\|[\s:|-]+\|$", lines[i + 1].strip()
        ):
            rows = [_split_row(stripped)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(_split_row(lines[i].strip()))
                i += 1
            _add_table(doc, rows)
            continue

        # --- 제목 ---
        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = re.sub(r"[*`]", "", m.group(2))
            h = doc.add_heading("", level=level)
            size = {1: 18, 2: 14, 3: 12, 4: 11}[level]
            _set_font(h.add_run(text), BODY_FONT, size, bold=True)
            i += 1
            continue

        # --- 수평선 ---
        if re.match(r"^-{3,}$", stripped):
            i += 1
            continue

        # --- 인용 ---
        if stripped.startswith(">"):
            quote: list[str] = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote.append(lines[i].strip().lstrip(">").strip())
                i += 1
            par = doc.add_paragraph()
            par.paragraph_format.left_indent = Pt(18)
            _add_inline(par, " ".join(q for q in quote if q), 9.5)
            for run in par.runs:
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            continue

        # --- 목록 ---
        m = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)$", line)
        if m:
            indent, marker, text = m.group(1), m.group(2), m.group(3)
            style = "List Number" if marker[0].isdigit() else "List Bullet"
            par = doc.add_paragraph(style=style)
            if len(indent) >= 2:
                par.paragraph_format.left_indent = Pt(36)
            _add_inline(par, text)
            i += 1
            continue

        # --- 빈 줄 ---
        if not stripped:
            i += 1
            continue

        # --- 일반 문단 ---
        par = doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _add_inline(par, stripped)
        i += 1

    doc.save(str(docx_path))


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print(__doc__)
        sys.exit(1)
    convert(Path(sys.argv[1]), Path(sys.argv[2]))
    print(f"생성: {sys.argv[2]}")
